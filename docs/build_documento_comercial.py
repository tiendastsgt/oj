# -*- coding: utf-8 -*-
"""
Genera SGED_Presentacion_Tecnica_v1.0.0.docx para presentación comercial SGED.
Fuentes de verdad: docs/general/plan_detallado.md, ROADMAP, STACK_TECNICO_ACTUALIZADO.md,
docs/qa/QA_ACCEPTANCE_REPORT.md, docs/diagramas/*.png.
Sin referencias internas (IA, Cursor, agentes, rutas locales) ni datos sensibles.
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = Path(__file__).resolve().parent
DIAGRAMAS = BASE / "diagramas"

def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    p.space_after = Pt(6)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    p.paragraph_format.space_after = Pt(3)
    return p

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- PORTADA / TÍTULO ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SGED — Sistema de Gestión de Expedientes Digitales")
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Documento de presentación comercial — v1.0.0 en producción")
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run("Enero 2026")
    doc.add_paragraph()
    doc.add_page_break()

    # --- 1. RESUMEN EJECUTIVO ---
    add_heading_custom(doc, "1. Resumen ejecutivo", 1)
    add_para(doc, "SGED es el Sistema de Gestión de Expedientes Digitales del Organismo Judicial, diseñado para centralizar la gestión de expedientes digitales, documentos multimedia y búsqueda unificada, con trazabilidad completa y control de acceso por roles.")
    add_para(doc, "El sistema se encuentra en producción desde enero de 2026 (versión 1.0.0), con todas las fases funcionales validadas y certificación QA como apto para producción. SGED integra en solo lectura los sistemas legacy SGTv1 y SGTv2, sin modificar sus datos.")
    add_para(doc, "Valor principal: seguridad (autenticación JWT, RBAC, auditoría inmutable), operación estable (objetivos de rendimiento cumplidos) y mantenibilidad (arquitectura modular y documentación técnica centralizada).")

    # --- 2. CONTEXTO Y VALOR ---
    add_heading_custom(doc, "2. Contexto y valor de negocio", 1)
    add_para(doc, "El Organismo Judicial opera con dos sistemas de gestión de tribunales: SGTv1 (histórico) y SGTv2 (activo). SGED aporta una capa unificada de expedientes digitales, documentos y búsqueda, sin sustituir dichos sistemas.")
    add_para(doc, "Beneficios clave:")
    add_para(doc, "• Gestión centralizada de expedientes digitales con CRUD completo y control por juzgado y rol.")
    add_para(doc, "• Carga y visualización de documentos multimedia (PDF, Word convertido a PDF, imágenes, audio y vídeo) con visores integrados.")
    add_para(doc, "• Búsqueda rápida y avanzada, con integración de resultados de SGTv1 y SGTv2 en una única interfaz.")
    add_para(doc, "• Trazabilidad: todas las operaciones se registran en auditoría inmutable.")
    add_para(doc, "• Seguridad: autenticación con JWT (8 h), contraseñas con BCrypt, bloqueo tras 5 intentos fallidos, y control de acceso por cuatro roles (Administrador, Secretario, Auxiliar, Consulta).")

    # --- 3. FUNCIONALIDADES v1.0.0 ---
    add_heading_custom(doc, "3. Funcionalidades incluidas (v1.0.0)", 1)
    add_para(doc, "Las siguientes capacidades están implementadas y validadas en producción:")
    add_para(doc, "• F01 — Gestión de expedientes: crear, consultar, editar y listar con filtros y paginación.")
    add_para(doc, "• F02 — Carga de documentos con validación de formato y tamaño (máx. 100 MB).")
    add_para(doc, "• F03 — Visor de documentos: PDF, Word (convertido a PDF), imágenes.")
    add_para(doc, "• F04 — Reproductor multimedia: audio y vídeo con controles HTML5.")
    add_para(doc, "• F05 — Búsqueda rápida por número de expediente.")
    add_para(doc, "• F06 — Búsqueda avanzada con filtros combinables e integración SGTv1/SGTv2.")
    add_para(doc, "• F07 — Integración SGT: consulta solo lectura a SGTv1 y SGTv2.")
    add_para(doc, "• F08 — Control de acceso: cuatro roles con permisos definidos.")
    add_para(doc, "• F09 — Auditoría: registro de todas las operaciones relevantes.")
    add_para(doc, "• F10 — Impresión de documentos desde el visor.")

    # --- 4. STACK TÉCNICO ---
    add_heading_custom(doc, "4. Stack técnico", 1)
    add_para(doc, "Frontend: Angular 21 LTS, TypeScript 5.7+, PrimeNG 21, Node.js 22 LTS. Compatible con Chrome, Edge y Firefox en versiones recientes.")
    add_para(doc, "Backend: Java 21 LTS, Spring Boot 3.5.x, Spring Security 6.5.x, Spring Data JPA, Hibernate 6.7.x, JJWT 0.12 para JWT. Build con Maven 3.10.x.")
    add_para(doc, "Base de datos: Oracle 21c (principal SGED). Conexiones de solo lectura a SGTv1 y SGTv2. Migraciones con Flyway.")
    add_para(doc, "Infraestructura: Docker, Docker Compose, NGINX como reverse proxy con HTTPS y endurecimiento de seguridad.")

    # --- 5. SEGURIDAD Y TRAZABILIDAD ---
    add_heading_custom(doc, "5. Seguridad y trazabilidad", 1)
    add_para(doc, "• Autenticación: JWT con expiración de 8 horas; revocación en cierre de sesión.")
    add_para(doc, "• Contraseñas: BCrypt (12 rounds); requisitos de complejidad (mínimo 8 caracteres, mayúscula, minúscula, número).")
    add_para(doc, "• Bloqueo de cuenta tras 5 intentos fallidos de inicio de sesión.")
    add_para(doc, "• Auditoría: registro asíncrono e inmutable de acciones (login, logout, cambios de contraseña, CRUD de expedientes y usuarios, consultas de auditoría, etc.).")
    add_para(doc, "• Endurecimiento: HTTPS/TLS, cabeceras de seguridad, rate limiting, y escaneos de seguridad (OWASP ZAP, CodeQL) con cero vulnerabilidades críticas o altas en el informe de QA.")

    # --- 6. QA Y ESTADO EN PRODUCCIÓN ---
    add_heading_custom(doc, "6. Aseguramiento de calidad y estado en producción", 1)
    add_para(doc, "El sistema ha sido validado en la Fase 7 de QA. Resultados resumidos:")
    add_para(doc, "• Pruebas E2E: flujos F1–F6 (login/RBAC, expedientes, documentos, búsqueda, gestión de usuarios, auditoría) ejecutados correctamente.")
    add_para(doc, "• Rendimiento: P95 de APIs < 2 s, P99 < 3 s; soporte validado para al menos 50 usuarios concurrentes (pruebas hasta 75 usuarios).")
    add_para(doc, "• Seguridad: 0 vulnerabilidades críticas o altas; controles de autenticación, autorización y auditoría verificados.")
    add_para(doc, "Conclusión del informe de aceptación: APTO PARA PRODUCCIÓN. La versión 1.0.0 está desplegada y operativa con despliegue controlado (rollout gradual).")

    # --- 7. DIAGRAMAS ---
    add_heading_custom(doc, "7. Diagramas de arquitectura y flujo", 1)
    add_para(doc, "A continuación se incluyen los diagramas vigentes que reflejan la arquitectura y los flujos del sistema.")
    labels = [
        ("diagram_1.png", "Arquitectura de capas: Frontend (Angular), Backend (Spring Boot), Base de datos (Oracle) y sistema de archivos."),
        ("diagram_2.png", "Flujo de interacción entre usuario, peticiones HTTP y persistencia."),
        ("diagram_3.png", "Flujo de autenticación: login, validación, JWT y auditoría."),
        ("diagram_4.png", "Pirámide de pruebas (unitarias, integración, E2E)."),
        ("diagram_5.png", "Despliegue: navegador, NGINX, backend y frontend, bases de datos y almacenamiento."),
        ("diagram_6.png", "Checklist post-despliegue: servicios, funcionalidad, seguridad y rendimiento."),
    ]
    for fname, caption in labels:
        path = DIAGRAMAS / fname
        if path.exists():
            add_para(doc, caption)
            try:
                doc.add_picture(str(path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                add_para(doc, f"[Diagrama: {fname}]")
            doc.add_paragraph()

    # --- 8. EQUIPO ---
    add_heading_custom(doc, "8. Equipo del proyecto", 1)
    add_para(doc, "• Amilcar Gil González — Arquitecto")
    add_para(doc, "• Emilio González González — Senior Developer")
    add_para(doc, "• Allan Gil González — Senior BE Developer")
    add_para(doc, "• Alejandro — PM (validación de calidad del producto)")
    add_para(doc, "La documentación técnica del proyecto se mantiene centralizada en el repositorio, con índice maestro y protocolo de vigencia.")

    # --- 9. CAPTURAS DE INTERFAZ DE USUARIO ---
    add_heading_custom(doc, "9. Capturas de interfaz de usuario", 1)
    add_para(doc, "Las siguientes capturas de pantalla corresponden a la interfaz de SGED v1.0.0. Donde aún no se dispone de imagen aprobada, se indica placeholder formal.")
    placeholders_ui = [
        "Captura pendiente: Inicio de sesión",
        "Captura pendiente: Listado de expedientes",
        "Captura pendiente: Detalle de expediente",
        "Captura pendiente: Visor documental",
        "Captura pendiente: Administración / roles",
    ]
    for ph in placeholders_ui:
        p = doc.add_paragraph()
        run = p.add_run(ph)
        run.italic = True
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)

    # --- Guardar ---
    out_dir = BASE / "general" / "presentacion"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "SGED_Presentacion_Tecnica_v1.0.0.docx"
    fallback = out_dir / "SGED_Presentacion_Tecnica_v1.0.0_PR.docx"
    try:
        doc.save(str(out))
        print(f"Documento generado: {out}")
    except PermissionError:
        doc.save(str(fallback))
        print(f"Documento generado (destino bloqueado): {fallback}")
    return str(out)

if __name__ == "__main__":
    main()
