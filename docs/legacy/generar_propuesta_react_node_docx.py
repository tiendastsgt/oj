# -*- coding: utf-8 -*-
"""Genera SGED_Presentacion_Tecnica_Propuesta_React_Node_vNext.docx. Una sola ejecución."""
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parent
DOCS = BASE / "docs"
DIAGRAMAS = DOCS / "diagramas"
PRESENTACION = DOCS / "general" / "presentacion"

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.space_after = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.space_after = Pt(4)
    return p

def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def main():
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ----- PORTADA -----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SGED – Sistema de Gestión de Expedientes Digitales")
    r.bold = True
    r.font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Propuesta técnica – Stack React 19.2.4 + Node.js 22.20.0")
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run("Fecha: Enero 2026")
    doc.add_paragraph()
    oj = doc.add_paragraph()
    oj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    oj.add_run("Organismo Judicial")
    doc.add_page_break()

    # ----- 1. RESUMEN EJECUTIVO -----
    h1(doc, "1. Resumen ejecutivo")
    h2(doc, "1.1 Problema que resuelve SGED")
    para(doc, "El Organismo Judicial opera con sistemas legacy de gestión de tribunales (SGTv1 y SGTv2). SGED centraliza la gestión de expedientes digitales, la carga y visualización de documentos multimedia, y la búsqueda unificada, con trazabilidad completa y control de acceso por roles. La propuesta técnica que se presenta define la implementación con un stack moderno (React + Node.js) que permita mantener y evolucionar el sistema con seguridad, auditoría y eficiencia operativa.")
    h2(doc, "1.2 Valor: trazabilidad, auditoría, seguridad, eficiencia, modernización")
    para(doc, "• Trazabilidad: todas las operaciones sobre expedientes y documentos quedan registradas en una tabla de auditoría inmutable.")
    para(doc, "• Auditoría: login/logout, cambios de contraseña, accesos a expedientes y documentos, búsquedas y acciones de administración se auditan con usuario, timestamp, acción y recurso afectado.")
    para(doc, "• Seguridad: autenticación JWT (8 h), política de contraseñas, bloqueo tras 5 intentos fallidos, revocación de tokens, rate limiting, cabeceras de seguridad y HTTPS obligatorio.")
    para(doc, "• Eficiencia: búsqueda rápida y avanzada, integración en solo lectura con SGTv1/SGTv2 (prioridad SGTv2), visores integrados para PDF e imágenes.")
    para(doc, "• Modernización: stack React 19.2.4 + TypeScript y Node.js 22 LTS, con arquitectura por features, API REST documentable y despliegue en contenedores (Docker + NGINX).")
    h2(doc, "1.3 Alcance funcional (alto nivel)")
    para(doc, "• Autenticación y sesión: inicio/cierre de sesión, cambio de contraseña, revocación de tokens.")
    para(doc, "• Gestión de expedientes: creación, consulta, edición y listado con filtros y control por juzgado y rol.")
    para(doc, "• Gestión documental y visores: carga de documentos (límite 100 MB), almacenamiento por año/mes, visores para PDF, Word (convertido a PDF), imágenes, audio y vídeo.")
    para(doc, "• Búsqueda rápida y avanzada: por número de expediente y filtros combinables, con integración SGTv1/SGTv2 (solo lectura).")
    para(doc, "• Administración: gestión de usuarios, asignación de roles, bloqueo/desbloqueo, reset de contraseña (solo ADMINISTRADOR).")
    para(doc, "• Auditoría: consulta de registros de auditoría con filtros (usuario, módulo, acción, rango de fechas).")
    doc.add_paragraph()

    # ----- 2. ARQUITECTURA PROPUESTA -----
    h1(doc, "2. Arquitectura propuesta (alto nivel)")
    para(doc, "• Frontend: React 19.2.4 con TypeScript 5.7+. Componentes funcionales y hooks. Arquitectura por features (auth, expedientes, documentos, búsqueda, administración). Routing con React Router. Estado: Context API o Zustand para estado global acotado; React Query para datos del servidor (cache e invalidación).")
    para(doc, "• Backend: Node.js 22 LTS (22.20.0). Framework recomendado: NestJS — arquitectura por módulos, TypeScript nativo, integración con Passport/JWT, validación y documentación OpenAPI; adecuado para aplicaciones enterprise y mantenibilidad. Alternativa: Express o Fastify si se prioriza ligereza y flexibilidad.")
    para(doc, "• Base de datos: Oracle 21c. Migraciones versionadas (por ejemplo Flyway o equivalente en Node). Conexiones de solo lectura a SGTv1 y SGTv2; prioridad SGTv2.")
    para(doc, "• Infraestructura: NGINX como reverse proxy, terminación TLS (HTTPS), cabeceras de seguridad y rate limiting. Aplicaciones en contenedores Docker.")
    para(doc, "• Observabilidad y logging: logs en formato JSON; métricas y health checks para monitoreo.")
    para(doc, "Diagrama lógico: se recomienda reutilizar los diagramas de arquitectura de capas y flujo ya existentes en el repositorio (docs/diagramas). Si en la revisión técnica no se reutilizan tal cual, incluir aquí un diagrama actualizado con React + Node + Oracle 21c + NGINX. [Placeholder: Diagrama pendiente de actualización si no se reutiliza el existente.]")
    doc.add_paragraph()

    # ----- 3. SEGURIDAD -----
    h1(doc, "3. Seguridad")
    para(doc, "• JWT: expiración 8 horas; firma con secreto gestionado por variables de entorno; revocación en cierre de sesión (lista de JTI revocados o equivalente).")
    para(doc, "• Política de contraseñas: mínimo 8 caracteres, al menos una mayúscula, una minúscula y un número; almacenamiento con hash (bcrypt u equivalente, factor de coste adecuado).")
    para(doc, "• Bloqueo de cuenta: tras 5 intentos fallidos de inicio de sesión, la cuenta se bloquea; desbloqueo solo por administrador.")
    para(doc, "• Revocación de tokens: al cerrar sesión el token se invalida; el backend rechaza tokens revocados.")
    para(doc, "• Rate limiting: límite de peticiones por IP en endpoints sensibles (login, API pública) para mitigar fuerza bruta y abuso.")
    para(doc, "• Cabeceras de seguridad: X-Content-Type-Options, X-Frame-Options, CSP según política del OJ, etc.")
    para(doc, "• HTTPS obligatorio: toda la comunicación usuario–sistema y entre servicios expuestos vía TLS.")
    para(doc, "• RBAC: cuatro roles — ADMINISTRADOR, SECRETARIO, AUXILIAR, CONSULTA — con permisos definidos por recurso y acción.")
    para(doc, "• Auditoría inmutable: tabla de auditoría con campos clave (fecha, usuario, IP, acción, módulo, recurso_id, valor_anterior/valor_nuevo si aplica, detalle); sin borrado ni edición de registros.")
    doc.add_paragraph()

    # ----- 4. AUDITORÍA Y TRAZABILIDAD -----
    h1(doc, "4. Auditoría y trazabilidad")
    para(doc, "Qué se audita: intentos de login (exitosos y fallidos), logout, cambio de contraseña, creación/edición/consulta de expedientes, carga y visualización/descarga de documentos, búsquedas relevantes, creación/actualización de usuarios, bloqueo/desbloqueo, reset de contraseña por administrador, y consultas a la propia auditoría.")
    para(doc, "Inmutabilidad y retención: los registros de auditoría no se modifican ni eliminan; la retención se define según normativa interna del OJ (conservación a largo plazo para fines de control y cumplimiento).")
    doc.add_paragraph()

    # ----- 5. MÓDULOS FUNCIONALES -----
    h1(doc, "5. Módulos funcionales (alto nivel)")
    para(doc, "• Autenticación y sesión: login, logout, cambio de contraseña, renovación/revocación de tokens.")
    para(doc, "• Gestión de expedientes: CRUD, listado paginado, filtros por juzgado/estado/tipo, control por rol.")
    para(doc, "• Gestión documental y visores: subida con validación de tipo y tamaño (máx. 100 MB), almacenamiento por año/mes, visores para PDF, Word (convertido a PDF), imágenes, reproductores audio/vídeo HTML5.")
    para(doc, "• Búsqueda rápida y avanzada: por número de expediente; búsqueda avanzada con filtros combinables; integración SGTv2 y SGTv1 en solo lectura (prioridad SGTv2).")
    para(doc, "• Integración SGTv2/v1: consultas read-only; sin escritura en sistemas legacy.")
    para(doc, "• Administración y consulta de auditoría: gestión de usuarios y roles (ADMINISTRADOR), bloqueo/desbloqueo, reset de contraseña; consulta de auditoría con filtros.")
    doc.add_paragraph()

    # ----- 6. ESTRATEGIA DE IMPLEMENTACIÓN -----
    h1(doc, "6. Estrategia de implementación y entregables (por fases)")
    para(doc, "• F0 — Base repo e infra: repositorios, Docker, NGINX base, CI (build y tests), documentación de convenciones.")
    para(doc, "• F1 — Autenticación y seguridad: login/logout, JWT, cambio de contraseña, lockout, RBAC base, auditoría de auth.")
    para(doc, "• F2 — Expedientes: modelo de datos, API CRUD, listado y filtros, UI React.")
    para(doc, "• F3 — Documentos y visor: subida, almacenamiento año/mes, visores PDF/imagen/audio/vídeo, conversión Word a PDF si aplica.")
    para(doc, "• F4 — Búsqueda + SGT: búsqueda rápida y avanzada, clientes read-only SGTv2/SGTv1, combinación de resultados.")
    para(doc, "• F5 — Administración y auditoría UI: gestión de usuarios, roles, bloqueo/reset; pantalla de consulta de auditoría con filtros.")
    para(doc, "• F6 — Hardening y rendimiento: cabeceras de seguridad, rate limiting, pruebas de carga, ajuste de índices y consultas.")
    para(doc, "• F7 — QA y go-live controlado: pruebas E2E, validación de criterios de aceptación, despliegue gradual (por ejemplo 10 % → 50 % → 100 %).")
    doc.add_paragraph()

    # ----- 7. REUTILIZACIÓN -----
    h1(doc, "7. Reutilización del sistema actual")
    h2(doc, "7.1 Qué se puede reutilizar sin rehacer")
    para(doc, "• Modelo de datos Oracle y scripts de migración (Flyway o equivalentes) si existen y son compatibles con Oracle 21c.")
    para(doc, "• Configuración de NGINX (hardening, headers, rate limiting) adaptada al nuevo stack.")
    para(doc, "• Documentación existente: plan funcional, criterios de QA, runbooks de operación, diagramas PlantUML/PNG (adaptando leyendas si cambia el stack).")
    para(doc, "• Criterios de QA y smoke tests (casos de prueba y criterios de aceptación reutilizables; implementación de tests en React/Node).")
    h2(doc, "7.2 Qué no se reutiliza directamente")
    para(doc, "• Frontend Angular: se reemplaza por la nueva implementación en React 19.2.4 + TypeScript.")
    para(doc, "• Backend Spring Boot (Java): se reemplaza por la nueva implementación en Node.js (NestJS o Express/Fastify).")
    para(doc, "La reutilización exacta dependerá de la revisión técnica del repositorio actual y de las decisiones de Arquitectura.")
    doc.add_paragraph()

    # ----- 8. EQUIPO -----
    h1(doc, "8. Equipo y roles")
    para(doc, "• Amilcar Gil González — Arquitecto")
    para(doc, "• Emilio González González — Senior Developer")
    para(doc, "• Allan Gil González — Senior BE Developer")
    para(doc, "• Alejandro — Project Manager (valida calidad del producto)")
    doc.add_paragraph()

    # ----- 9. CIERRE Y PRÓXIMOS PASOS -----
    h1(doc, "9. Cierre y próximos pasos")
    para(doc, "Para iniciar la implementación, el OJ deberá disponer de:")
    para(doc, "• Infraestructura: acceso a entornos (desarrollo, QA, producción) con Oracle 21c, o definición de instancias y credenciales gestionadas de forma segura.")
    para(doc, "• Accesos: cuentas y permisos para repositorios, CI/CD y despliegue; acceso de solo lectura a SGTv1/SGTv2 según acuerdos.")
    para(doc, "• Ambientes: Docker y/o servidores con Node.js 22 LTS y NGINX; certificados TLS para HTTPS.")
    para(doc, "• Datos de prueba: conjunto anonimizado o de prueba para expedientes y documentos, y usuarios de prueba por rol, para validar flujos y auditoría.")
    para(doc, "Con estos elementos, el equipo puede alinear la F0 y arrancar la F1 (autenticación y seguridad) según la estrategia de fases descrita.")

    # ----- GUARDAR -----
    PRESENTACION.mkdir(parents=True, exist_ok=True)
    out = PRESENTACION / "SGED_Presentacion_Tecnica_Propuesta_React_Node_vNext.docx"
    doc.save(str(out))
    print(f"Generado: {out}")
    return str(out)

if __name__ == "__main__":
    main()
