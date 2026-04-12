import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_formatted_text(paragraph, text):
    """
    Busca texto entre ** ** y lo pone en negrita.
    """
    # Expresión regular para encontrar **texto**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Es negrita, quitamos los asteriscos
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Texto normal
            paragraph.add_run(part)

def is_ascii_art(line):
    box_chars = set("┌┐└┘├┤┬┴┼─│═║╔╗╚╝╠╣╦╩╬")
    return any(char in box_chars for char in line)

def is_table_row(line):
    # Detecta tablas con barras vertical OR tablas separadas por tabulaciones
    has_pipes = line.strip().startswith('|') or ('|' in line and len(line.split('|')) > 2)
    has_tabs = '\t' in line and len(line.split('\t')) > 1
    return has_pipes or has_tabs

def process_table_row(line):
    # Limpia la línea y devuelve una lista de celdas
    if '|' in line:
        return [c.strip() for c in line.split('|') if c.strip() != '']
    else:
        return [c.strip() for c in line.split('\t') if c.strip() != '']

def create_professional_doc(input_file, output_file):
    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Lectura robusta del archivo
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except UnicodeDecodeError:
        print("Detectada codificación antigua, leyendo en modo compatibilidad...")
        with open(input_file, 'r', encoding='latin-1') as f:
            lines = f.readlines()

    buffer_table = []
    in_code_block = False # Para detectar ``` bloque de codigo ```

    for line in lines:
        stripped = line.strip()
        raw_line = line.rstrip() # Mantiene espacios a la derecha para ASCII art

        # --- 1. DETECCIÓN DE BLOQUES DE CÓDIGO (```) ---
        if stripped.startswith('```'):
            in_code_block = not in_code_block # Entrar o salir del modo código
            continue # No imprimir la línea de ```
        
        if in_code_block or (is_ascii_art(line) and not in_code_block):
            # Imprimir como código / dibujo (fuente Consolas)
            # Primero vaciamos buffer de tablas si existe
            if buffer_table:
                create_word_table(doc, buffer_table)
                buffer_table = []
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(raw_line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            continue

        # --- 2. DETECCIÓN DE TABLAS ---
        # Si la línea parece tabla, la guardamos en el buffer
        if is_table_row(line) and not stripped.startswith('#'):
            buffer_table.append(line)
            continue
        
        # Si ya no es tabla pero tenemos datos guardados, CREAMOS LA TABLA
        if buffer_table:
            create_word_table(doc, buffer_table)
            buffer_table = [] # Limpiar buffer
            if stripped == "": continue # Ignorar línea vacía después de tabla

        # --- 3. FORMATO DE TEXTO (Títulos, Listas, etc) ---
        
        # Línea vacía
        if stripped == "":
            continue

        # Línea divisoria (---)
        if set(stripped) == {'-'}:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            # Borde inferior para simular línea
            p_element = p._p
            pPr = p_element.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pbdr.append(bottom)
            pPr.append(pbdr)
            continue

        # Títulos (#, ##, ###)
        if stripped.startswith('#'):
            level = stripped.count('#', 0, 10) # Contar cuántos # hay al inicio
            text = stripped.lstrip('#').strip()
            # Word soporta hasta nivel 9, aseguramos no pasarnos
            level = min(level, 9)
            if level > 0:
                h = doc.add_heading(level=level)
                add_formatted_text(h, text)
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped)
            continue

        # Citas (>)
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20) # Sangría
            run = p.add_run("│ " + text) # Barra visual
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80) # Gris oscuro
            continue

        # Listas (-)
        if stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        # Párrafo Normal (con detección de negrita)
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)

    # Limpieza final por si quedó tabla pendiente
    if buffer_table:
        create_word_table(doc, buffer_table)

    doc.save(output_file)
    print(f"¡Éxito! Documento generado: {output_file}")

def create_word_table(doc, rows_data):
    # Filtrar líneas que sean solo separadores Markdown (ej: |---|---|)
    clean_rows = []
    for row in rows_data:
        # Si la línea tiene letras o números, es válida. Si solo tiene guiones/barras, es separador.
        if any(c.isalnum() for c in row):
            clean_rows.append(row)
    
    if not clean_rows: return

    # Analizar columnas basado en la primera fila
    headers = process_table_row(clean_rows[0])
    cols = len(headers)
    
    if cols == 0: return

    table = doc.add_table(rows=len(clean_rows), cols=cols)
    table.style = 'Table Grid'
    table.autofit = True

    for i, row_text in enumerate(clean_rows):
        cells_data = process_table_row(row_text)
        for j, cell_text in enumerate(cells_data):
            if j < cols:
                cell = table.cell(i, j)
                # Limpiar párrafo por defecto
                cell.text = "" 
                p = cell.paragraphs[0]
                # Usar la función de formato para mantener negritas dentro de la tabla
                add_formatted_text(p, cell_text)
                
                # Si es la primera fila (encabezado), forzar negrita y fondo gris claro
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                    # Sombreado (XML complejo omitido para simplicidad, pero la negrita basta)

# Ejecutar
try:
    create_professional_doc('entrada.txt', 'Documento_Final.docx')
except FileNotFoundError:
    print("Error: Crea el archivo 'entrada.txt' con tu texto antes de ejecutar.")